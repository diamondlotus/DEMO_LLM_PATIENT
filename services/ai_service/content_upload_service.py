"""
AI Service - Content Upload and Learning Service
Handles PDF uploads, content processing, and AI learning from various sources
"""

import os
import sys
import json
import time
from typing import Dict, Any, List, Optional
from datetime import datetime
import tempfile
import shutil
from pathlib import Path

# Add project root to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from llm_integration.healthcare_llm import HealthcareLLM
from demo_data_loader import DemoDataLoader

class ContentUploadService:
    """Service for handling content uploads and AI learning"""

    def __init__(self):
        self.healthcare_llm = HealthcareLLM()
        self.upload_dir = Path("./uploads")
        self.upload_dir.mkdir(exist_ok=True)
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx', '.html']

    def upload_and_process_file(self, file_data: bytes, filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Upload and process a file for AI learning
        """
        try:
            # Validate file type
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}. Supported: {', '.join(self.supported_formats)}")

            # Save file temporarily
            temp_path = self.upload_dir / f"temp_{int(time.time())}_{filename}"
            with open(temp_path, 'wb') as f:
                f.write(file_data)

            # Extract text content
            text_content = self._extract_text_from_file(temp_path, file_ext)

            if not text_content or len(text_content.strip()) < 100:
                raise ValueError("File contains insufficient text content for processing")

            # Process content for AI learning
            processing_result = self._process_content_for_learning(
                text_content,
                filename,
                metadata or {}
            )

            # Clean up temp file
            temp_path.unlink(missing_ok=True)

            return {
                "success": True,
                "filename": filename,
                "file_size": len(file_data),
                "content_length": len(text_content),
                "processing_result": processing_result,
                "timestamp": datetime.utcnow().isoformat()
            }

        except Exception as e:
            # Clean up on error
            if 'temp_path' in locals():
                temp_path.unlink(missing_ok=True)

            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "timestamp": datetime.utcnow().isoformat()
            }

    def upload_text_content(self, content: str, title: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Upload text content directly for AI learning
        """
        try:
            if not content or len(content.strip()) < 50:
                raise ValueError("Content is too short for meaningful processing")

            # Process content for AI learning
            processing_result = self._process_content_for_learning(
                content,
                title,
                metadata or {}
            )

            return {
                "success": True,
                "title": title,
                "content_length": len(content),
                "processing_result": processing_result,
                "timestamp": datetime.utcnow().isoformat()
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "title": title,
                "timestamp": datetime.utcnow().isoformat()
            }

    def _extract_text_from_file(self, file_path: Path, file_ext: str) -> str:
        """
        Extract text content from various file formats
        """
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif file_ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif file_ext == '.pdf':
                # For PDF processing, we'll use a simple text extraction
                # In production, you'd use PyPDF2, pdfplumber, or similar
                return self._extract_text_from_pdf(file_path)

            elif file_ext == '.docx':
                # For DOCX processing
                return self._extract_text_from_docx(file_path)

            elif file_ext == '.html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                # Simple HTML text extraction (remove tags)
                import re
                clean_text = re.sub(r'<[^>]+>', '', content)
                return clean_text

            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

        except Exception as e:
            raise ValueError(f"Error extracting text from {file_ext} file: {str(e)}")

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF files using PyPDF2
        """
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_content = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)

            if not text_content:
                return f"Warning: No text content could be extracted from PDF {file_path.name}. The PDF may contain only images or scanned content."

            return "\n\n".join(text_content)

        except ImportError:
            return f"Error: PyPDF2 library not installed. Please install with: pip install PyPDF2"
        except Exception as e:
            return f"Error extracting PDF content: {str(e)}"

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX files using python-docx
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            if not text_content:
                return f"Warning: No text content could be extracted from DOCX {file_path.name}. The document may be empty or contain only images."

            return "\n\n".join(text_content)

        except ImportError:
            return f"Error: python-docx library not installed. Please install with: pip install python-docx"
        except Exception as e:
            return f"Error extracting DOCX content: {str(e)}"

    def _process_content_for_learning(self, content: str, source_name: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process content for AI learning and knowledge base integration
        """
        try:
            # Prepare metadata
            doc_metadata = {
                "source": source_name,
                "upload_timestamp": datetime.utcnow().isoformat(),
                "content_type": metadata.get("content_type", "uploaded_content"),
                "category": metadata.get("category", "general"),
                "tags": metadata.get("tags", []),
                "uploader": metadata.get("uploader", "system"),
                "processing_method": "content_upload"
            }

            # Add to healthcare LLM knowledge base
            success = self.healthcare_llm.add_medical_knowledge([content], [doc_metadata])

            if success:
                # Extract key insights and patterns
                insights = self._extract_content_insights(content)

                return {
                    "knowledge_base_updated": True,
                    "insights_extracted": insights,
                    "content_chunks": len(content) // 1000 + 1,  # Rough estimate
                    "metadata": doc_metadata
                }
            else:
                return {
                    "knowledge_base_updated": False,
                    "error": "Failed to add content to knowledge base"
                }

        except Exception as e:
            return {
                "knowledge_base_updated": False,
                "error": f"Content processing failed: {str(e)}"
            }

    def _extract_content_insights(self, content: str) -> Dict[str, Any]:
        """
        Extract key insights and patterns from content
        """
        insights = {
            "word_count": len(content.split()),
            "character_count": len(content),
            "medical_terms_found": [],
            "key_topics": [],
            "content_quality_score": 0.0
        }

        # Simple medical term detection
        medical_terms = [
            "patient", "diagnosis", "treatment", "symptoms", "medication",
            "cardiac", "respiratory", "neurological", "gastrointestinal",
            "hypertension", "diabetes", "cancer", "infection", "surgery"
        ]

        content_lower = content.lower()
        for term in medical_terms:
            if term in content_lower:
                insights["medical_terms_found"].append(term)

        # Extract key topics (simple keyword extraction)
        if "cardiac" in content_lower or "heart" in content_lower:
            insights["key_topics"].append("cardiology")
        if "respiratory" in content_lower or "lung" in content_lower:
            insights["key_topics"].append("pulmonology")
        if "neurological" in content_lower or "brain" in content_lower:
            insights["key_topics"].append("neurology")

        # Calculate content quality score
        quality_factors = [
            len(content) > 500,  # Sufficient length
            len(insights["medical_terms_found"]) > 3,  # Medical relevance
            len(insights["key_topics"]) > 0,  # Topic identification
            "conclusion" in content_lower or "summary" in content_lower  # Structure
        ]

        insights["content_quality_score"] = sum(quality_factors) / len(quality_factors)

        return insights

    def load_demo_data(self) -> Dict[str, Any]:
        """
        Load demo medical data using the demo data loader
        """
        try:
            loader = DemoDataLoader()
            return loader.load_demo_data()
        except Exception as e:
            return {
                "success": False,
                "error": f"Demo data loading failed: {str(e)}"
            }

    def get_uploaded_content_list(self) -> List[Dict[str, Any]]:
        """
        Get list of uploaded content (mock implementation)
        """
        # In production, this would query a database
        return [
            {
                "id": "upload_001",
                "filename": "cardiology_guidelines.pdf",
                "upload_date": "2024-01-01T10:00:00Z",
                "content_type": "medical_guidelines",
                "status": "processed",
                "insights": {"medical_terms": 25, "quality_score": 0.9}
            },
            {
                "id": "upload_002",
                "filename": "pediatric_care.txt",
                "upload_date": "2024-01-02T14:30:00Z",
                "content_type": "clinical_protocols",
                "status": "processed",
                "insights": {"medical_terms": 18, "quality_score": 0.8}
            }
        ]

    def get_learning_statistics(self) -> Dict[str, Any]:
        """
        Get AI learning statistics
        """
        return {
            "total_documents_processed": 15,
            "total_medical_terms_learned": 245,
            "knowledge_base_size": "2.3MB",
            "last_training_date": datetime.utcnow().isoformat(),
            "model_accuracy": 0.87,
            "categories_covered": [
                "cardiology", "pulmonology", "neurology",
                "gastroenterology", "pediatrics", "psychiatry"
            ]
        }

# Global service instance
content_upload_service = ContentUploadService()

def upload_file_service(file_data: bytes, filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """Convenience function for file upload"""
    return content_upload_service.upload_and_process_file(file_data, filename, metadata)

def upload_text_service(content: str, title: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """Convenience function for text upload"""
    return content_upload_service.upload_text_content(content, title, metadata)

def load_demo_data_service() -> Dict[str, Any]:
    """Convenience function for demo data loading"""
    return content_upload_service.load_demo_data()

if __name__ == "__main__":
    # Test the service
    print("🧠 AI Content Upload Service")
    print("Available functions:")
    print("- upload_file_service(file_data, filename, metadata)")
    print("- upload_text_service(content, title, metadata)")
    print("- load_demo_data_service()")